
import os
import sys
import unittest
from docx import Document

# Setup paths
PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))
sys.path.append(PROJECT_ROOT)
SAMPLE_RESUME = os.path.join(PROJECT_ROOT, 'sample_resume.docx')

def create_sample_resume():
    if os.path.exists(SAMPLE_RESUME):
        print(f"[{SAMPLE_RESUME}] already exists.")
        return

    print("Creating sample resume...")
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Email: john.doe@example.com | Phone: 123-456-7890')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, Flask, SQL, Machine Learning')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Engineer at Tech Corp (2020-Present)')
    doc.add_paragraph('Developed web applications using Flask and React.')
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.S. Computer Science, University of Examples')
    doc.add_heading('Links', level=1)
    doc.add_paragraph('GitHub: https://github.com/johndoe')
    doc.add_paragraph('LinkedIn: https://linkedin.com/in/johndoe')
    
    doc.save(SAMPLE_RESUME)
    print("Sample resume created successfully.")

if __name__ == "__main__":
    create_sample_resume()
    
    # Run tests
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()
    
    # Discover tests in 'tests' directory
    # Alternatively, load specific tests we know about
    try:
        suite.addTests(loader.discover('tests', pattern='verify_*.py'))
        print("Discovered verification tests.")
    except Exception as e:
        print(f"Error discovering tests: {e}")

    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    
    if result.wasSuccessful():
        print("\nAll integrity checks PASSED.")
        sys.exit(0)
    else:
        print("\nSome integrity checks FAILED.")
        sys.exit(1)
